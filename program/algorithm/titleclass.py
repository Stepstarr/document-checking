# -*- coding: utf-8 -*-
# @Time    : 2024/4/16 15:22
# @Author  : Stepstar
# @FileName: Title.py
# @Software: PyCharm
from docx import Document
from win32com.client import constants
from flask import Flask, Response

from algorithm.other_function.compare_attributes import compare_attributes
from baseclass import BaseClass
from flask import g
from format_config import title_format
import re
import time
from other_function.regex_expression import t_number_column_item, t_letter_column_item, get_text_pre
from detection_method.title_identify import title_identify, title_


class TitleClass(BaseClass):
    def __init__(self,doc,start = 1,end = None):
        super().__init__(doc)
        self.title_dict = {'标题 1': [], '标题 2': [], '标题 3': [], '标题 4': [], '标题 5': []}
        self.reference = {'字母列项': [], '数字列项': []}
        self.start = start
        if end is not None:
            self.end = end
        else:
            self.end = len(doc.Paragraphs)
        self.end = len(doc.Paragraphs)
        self.correct_number = {}
        self.get_titles()
        self.get_reference()

    def get_range(self):
        '''
        获取检测范围，附录前，标题后
        :return: 检测范围的开始和结束Page，Page和prg之间关系？能得到Page开始和结束prg吗？
        '''
    def get_titles(self):
        for index,p in enumerate(self.paragraphs):
            flag, title_layer = title_identify(p)
            if flag:
                self.title_dict[title_layer].append(index+1)
        self.correct_number = title_(self.title_dict)

    def get_reference(self):
        l = 0
        n = 0
        for index, p in enumerate(self.paragraphs):
            if t_letter_column_item(p.Range.ListFormat.ListString):
                self.reference['字母列项'].append(index+1)
                right_letter = chr(l + ord('a'))
                l += 1
                n = 0 #更新数字列项
                if p.Range.ListFormat.ListString != f'{right_letter})':
                    self.errors.append( {'错误段落':index+1,
                                        '段落Range':None,
                                         '类型':'correct',
                                        '修改属性': {'p.Range.ListFormat.ListString':f'{right_letter})'},
                                        '批注':f'字母列项错误，当前值为{p.Range.ListFormat.ListString}，正确值为{right_letter})'
                                     })
            elif t_number_column_item(p.Range.ListFormat.ListString):
                skip_increment = False  # 设置一个标志变量，初始为False
                for k in self.title_dict.keys():
                    if index + 1 in self.title_dict[k]:
                        skip_increment = True  # 如果条件满足，则设置标志为True
                        break  # 并且跳出for循环

                if not skip_increment:  # 检查标志，如果没有设置跳过，则增加n
                    n += 1
                    self.reference['数字列项'].append(index + 1)
                    right_letter = str(n)
                    if p.Range.ListFormat.ListString != f'{right_letter})':
                        self.errors.append({'错误段落': index + 1,
                                            '段落Range': None,
                                            '类型': 'correct',
                                            '修改属性': {'p.Range.ListFormat.ListString': f'{right_letter})'},
                                            '批注': f'数字列项错误，当前值为{p.Range.ListFormat.ListString}，正确值为{right_letter})'
                                            })
            elif t_letter_column_item(p.Range.Text):
                self.reference['字母列项'].append(index + 1)
                right_letter = chr(l + ord('a'))
                l += 1
                n = 0  # 更新数字列项
                group_0, group_1 = t_letter_column_item(p.Range.Text)  # 具体是否要检查再议
                self.errors.append({'错误段落': index + 1,
                                    '段落Range': None,
                                    '类型': 'correct',
                                    '修改属性': {'p.Range.ListFormat.ListString': f'{right_letter})'},
                                    '批注': f'字母列项应当使用自动编号'
                                    })
            elif t_number_column_item(p.Range.Text):
                skip_increment = False  # 设置一个标志变量，初始为False
                for k in self.title_dict.keys():
                    if index + 1 in self.title_dict[k]:
                        skip_increment = True  # 如果条件满足，则设置标志为True
                        break  # 并且跳出for循环

                if not skip_increment:  # 检查标志，如果没有设置跳过，则增加n
                    n += 1
                    self.reference['数字列项'].append(index + 1)
                    right_letter = str(n) + ')'
                    group_0, group_1 = t_number_column_item(p.Range.Text)
                    self.errors.append({'错误段落': index + 1,
                                    '段落Range': None,
                                    '类型': 'correct',
                                    '修改属性': {'p.Range.ListFormat.ListString': f'{right_letter})'},
                                    '批注': f'数字列项应当使用自动编号'
                                    })
    def check(self):
        def check_content():
            for index,value in self.correct_number.items():
                p = self.paragraphs(index)
                if str(p.Range.ListFormat.ListString) == '':
                    number, flag = get_text_pre(p.Range.Text)
                    if str(number) != str(value) + '  ':
                        clean_number = re.sub(r'\s+', '', str(number))
                        if clean_number == str(value):
                            if flag:
                                self.errors.append({'错误段落': index,
                                                    '段落Range': None,
                                                    '类型': 'warning',
                                                    '修改属性':None,
                                                    '批注': f'标题序号后无文本,需检查'
                                                    })
                            else:
                                self.errors.append({'错误段落': index,
                                                    '段落Range': None,
                                                    '类型': 'correct',
                                                    '修改属性': {'p.Range.ListFormat.ListString':{value}},
                                                    '批注': f'标题缩进或空格错误'
                                                    })
                        else:
                            self.errors.append({'错误段落': index,
                                                '段落Range': None,
                                                '类型': 'correct',
                                                '修改属性': {'p.Range.ListFormat.ListString': {value}},
                                                '批注': f'标题序号错误:当前值 {clean_number}，正确值应为 {str(value)}'
                                                })

                elif str(p.Range.ListFormat.ListString) != str(value):
                    self.errors.append({'错误段落': index,
                                        '段落Range': None,
                                        '类型': 'correct',
                                        '修改属性': {'p.Range.ListFormat.ListString': {value}},
                                        '批注': f'标题序号错误:当前值 {p.Range.ListFormat.ListString}，正确值应为 {str(value)}'
                                        })

                elif str(p.Range.ListFormat.ListString) == str(value):
                    space, flag = get_text_pre(p.Range.Text) # 得到文本前的字符
                    if flag:
                        self.errors.append({'错误段落': index,
                                                '段落Range': None,
                                                '类型': 'warning',
                                                '修改属性': None,
                                                '批注': f'标题序号后无文本,需检查'
                                                })

                    elif space != '  ':
                        self.errors.append({'错误段落': index,
                                            '段落Range': None,
                                            '类型': 'correct',
                                            '修改属性': {'p.Range.ListFormat.ListString': {value}},
                                            '批注': f'标题缩进或空格错误'
                                            })

        def check_style():
            for layer, indexs in self.title_dict.items():
                right_format = title_format.get(layer, {})
                for index in indexs:
                    p = self.paragraphs(index)
                    e = compare_attributes(index,p,right_format)
                    if e:
                        self.errors.append(e)
            for layer, indexs in self.reference.items():
                right_format =  title_format.get(layer, {})
                for index in indexs:
                    p = self.paragraphs(index)
                    e = compare_attributes(index,p,right_format)
                    if e:
                        self.errors.append(e)
        check_style()
        check_content()

    def add(self):
        for e in self.errors:
            p_range = self.paragraphs(e['错误段落']).Range
            comment_text = e['批注']
            comment = p_range.Comments.Add(Range=p_range, Text=comment_text)
            comment.Author = "HFUT checker"
            # 保存并关闭文档
            self.doc.Save()


    def correct(self):
        pass

        # 然后执行额外的检查
        # for attr, correct_value in additional_checks.items():
        #     actual_value = getattr(self.style, attr, None)
        #     if actual_value != correct_value:
        #         yield f"data: 新增属性 {attr} 错误: 当前值 {actual_value}，正确值应为 {correct_value}\n\n"
        #         time.sleep(1)  # 根据需要调整延迟


import win32com.client as win32
if __name__=='__main__':
    doc_app = win32.Dispatch("Kwps.Application")
    path = r'D:\Mo\项目\航发-标准化检测\program\static\航发标题层级测试.docx'
    doc_app.Visible = 1  # 设置应用程序可见
    doc = doc_app.Documents.Open(path, ReadOnly=True)  # 打开文档
    title_test = TitleClass(doc)
    title_test.check()
    title_test.add()

